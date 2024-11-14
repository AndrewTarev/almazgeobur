import os
import pathlib

from docx import Document
from fastapi import HTTPException
from openai import OpenAI

from src.core.config import settings
from src.core.utils.logging_config import my_logger


class GenerateReport:
    def __init__(self, date, total_revenue, top_products, categories):
        self.date = date
        self.total_revenue = total_revenue
        self.top_products = top_products
        self.categories = categories

    def _generate_report_content(self) -> str:
        """
        Генерирует аналитический отчет на основе предоставленных данных о продажах.

        :param self.date: Дата отчета.
        :param self.total_revenue: Общая выручка за период.
        :param self.top_products: Список топ-3 товаров по продажам.
        :param self.categories: Список категорий с их долей в продажах.
        :return: Сгенерированный текст аналитического отчета.
        """

        prompt = f"""
        Проанализируй данные о продажах за {self.date}:
        1. Общая выручка: {self.total_revenue}
        2. Топ-3 товара по продажам: {self.top_products}
        3. Распределение по категориям: {self.categories}
    
        Составь краткий аналитический отчет с выводами и рекомендациями.
        """
        try:
            client = OpenAI(
                api_key=settings.openapi.openai_api_key,
                base_url=settings.openapi.openai_base_url,
            )
            completion = client.chat.completions.create(
                model=settings.openapi.openai_gpt_model,
                messages=[{"role": "user", "content": prompt}],
            )
            report_content = completion.choices[0].message.content

        except Exception as e:
            my_logger.error(f"Error in OpenAI completion: {str(e)}")
            raise HTTPException(status_code=404, detail=str(e))

        my_logger.info("AI query completed successfully.")
        return report_content

    def _format_and_save_report(self, report_content):
        """
        Форматирует данные и сохраняет их в файл отчета в формате Word.

        :param self.date: Дата отчета.
        :param self.total_revenue: Общая выручка за период.
        :param self.top_products: Список топ-3 товаров по продажам.
        :param self.categories: Список категорий с их долей в продажах.
        :param report_content: Контент аналитического отчета.
        :return: Путь к сохраненному файлу отчета.
        """

        # Создаем документ Word
        doc = Document()

        # Добавляем заголовок
        doc.add_heading("Отчет о продажах", level=1)
        doc.add_paragraph(f"Дата отчета: {self.date}")

        # Добавляем разделы
        doc.add_heading("1. Общая выручка", level=2)
        doc.add_paragraph(f"Выручка за период: {self.total_revenue} руб.")

        doc.add_heading("2. Топ-3 товара по продажам", level=2)
        doc.add_paragraph(self.top_products, style="List Bullet")

        doc.add_heading("3. Распределение по категориям", level=2)
        doc.add_paragraph(self.categories, style="List Bullet")

        doc.add_heading("4. Аналитические выводы", level=2)
        doc.add_paragraph(report_content)

        # Указываем директорию для сохранения отчетов
        reports_dir = pathlib.Path(__file__).resolve().parent.parent.parent / "reports"
        if not os.path.exists(reports_dir):
            os.makedirs(reports_dir)
            my_logger.info(f"Created reports directory: {reports_dir}")

        # Создаем имя файла для отчета
        report_filename = os.path.join(reports_dir, f"report_sales_{self.date}.docx")

        # Сохраняем документ
        doc.save(report_filename)
        my_logger.info(f"Report saved successfully at {report_filename}")

        return report_filename

    def generate_and_save_report(self):
        """
        Генерирует и сохраняет отчет о продажах.

        :return: Строка с подтверждением сохранения отчета.
        """
        my_logger.info("Starting report generation process.")
        content = self._generate_report_content()

        report_filename = self._format_and_save_report(content)
        return f"Отчет сохранен в файле {report_filename}"

    def __call__(self):
        self.generate_and_save_report()
